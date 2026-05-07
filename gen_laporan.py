from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = r"d:\MAGANG MAHIRA TOUR\KERJA MAHIRA TOUR\LAPORAN APRIL"
os.makedirs(OUT, exist_ok=True)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('LAPORAN BULANAN KARYAWAN')
r.bold = True; r.font.size = Pt(16); r.font.name = 'Times New Roman'

for label, val in [('NAMA KARYAWAN','KHOIRUL GUNAWAN'),('JABATAN','IT'),('PERIODE','APRIL 2026')]:
    p = doc.add_paragraph()
    r = p.add_run(f'{label}\t\t\t: {val}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.add_paragraph()

# Table
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, txt in enumerate(['KEGIATAN','CAPAIAN']):
    tbl.rows[0].cells[i].text = txt
    for p in tbl.rows[0].cells[i].paragraphs:
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(11)

data = [
    ('Finalisasi Aplikasi Sautul Haram\n(Radio Muthawwif)',
     'Implementasi fitur notifikasi sesi baru berhasil dilakukan.\n'
     'Konten doa, shalawat, dan surah (termasuk Surah Yasin) berhasil ditambahkan.\n'
     'Setup offline storage dan halaman riwayat sesi muthawwif berhasil.\n'
     'Perbaikan bug tampilan beranda, login muthawwif, alur logika sesi, dan UI/UX.\n'
     'Testing di device fisik Android berhasil, APK release tersedia (~100MB).\n'
     'Redesain komponen UI/UX, 7+ improvement, konfigurasi standar keamanan.\n'
     'Optimasi performa, font arab, tombol navigasi, dan implementasi corak islami.'),
    ('Maintenance & Pengembangan\nWebsite Mahira Tour',
     'Maintenance rutin website berjalan lancar tanpa kendala.\n'
     'Update dokumentasi dan konten website secara berkala.\n'
     'Publish artikel: tour Singapura-Malaysia, testimoni jamaah, lowongan finance.\n'
     'Pembuatan PIC website agar bisa dipelajari oleh staff lain.\n'
     'Pembayaran dan pengecekan server berhasil dilakukan.'),
    ('Perancangan Proyek\nSistem Laporan Terintegrasi',
     'Riset mendalam proyek terakhir magang, evaluasi opsi proyek.\n'
     'Blueprint dan flowchart sistem berhasil dirancang.\n'
     'Setup awal proyek dan database (Next.js + Supabase) dimulai.\n'
     'Konsep alur kerja harian, mingguan, dan bulanan sudah dirancang.'),
    ('Tugas Operasional',
     'Mengambil 2 spanduk Ust. Nadirman dari mersi dan dibawa ke kantor.\n'
     'Mengambil brosur dari mersi, packing, dan kirim via JNT.'),
]

for k, c in data:
    row = tbl.add_row().cells
    row[0].text = k; row[1].text = c
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

for row in tbl.rows:
    row.cells[0].width = Cm(5.5); row.cells[1].width = Cm(12)

doc.add_paragraph(); doc.add_paragraph()

p = doc.add_paragraph(); r = p.add_run('Sungai penuh,  1 Mei 2026'); r.font.name = 'Times New Roman'
p = doc.add_paragraph(); r = p.add_run('Disetujui oleh:\t\t\t\t\t\t\t\t\tHormat Saya,'); r.font.name = 'Times New Roman'
for _ in range(3): doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run('Khilal, S.Th.I\t\t\tNadirman\t\t\t\t\t\t\tKhoirul Gunawan'); r.font.name = 'Times New Roman'
p = doc.add_paragraph(); r = p.add_run('Direktur Utama\t\tKomisaris Utama\t\t\t\t\t\tIT'); r.font.name = 'Times New Roman'

fp = os.path.join(OUT, 'LAPORAN BULAN APRIL.docx')
doc.save(fp)
print(f'[OK] {fp}')
