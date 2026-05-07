from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUT = r"d:\MAGANG MAHIRA TOUR\KERJA MAHIRA TOUR\LAPORAN APRIL"
doc = Document()
s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(11)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.name = 'Times New Roman'

def para(text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# === COVER ===
doc.add_paragraph()
para('BLUEPRINT', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
para('SISTEM LAPORAN TERINTEGRASI', bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
para('MAHIRA TOUR', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Versi: 1.0', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Tanggal: 28 April 2026', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Domain: laporan.mahiratour.id', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Timeline: 4 Minggu', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para('Disusun oleh:', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
para('Khoirul Gunawan - Staff IT / FullStack Developer', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# === 1. PENDAHULUAN ===
heading('1. Pendahuluan')
para('Sistem Laporan Terintegrasi adalah sebuah web application yang dirancang untuk mendigitalkan proses pelaporan kerja harian, arsip dokumen, dan komunikasi internal di Mahira Tour. Sistem ini dapat diakses melalui HP (mobile-first) maupun laptop oleh seluruh karyawan (~15-20 orang) dan pimpinan (2 orang direksi).')

heading('1.1 Latar Belakang Masalah', level=2)
bullet('Laporan kerja harian masih dilakukan secara manual melalui WhatsApp')
bullet('Tidak ada sistem terpusat untuk menyimpan dan mengelola arsip dokumen per divisi')
bullet('Pimpinan kesulitan memantau progres kerja seluruh divisi secara real-time')
bullet('Tidak ada sistem reminder otomatis untuk pengumpulan laporan')
bullet('Data laporan harian tersebar dan sulit dicari kembali')

heading('1.2 Tujuan Sistem', level=2)
bullet('Menyediakan platform digital untuk input dan submit laporan kerja harian')
bullet('Memudahkan pimpinan memonitor kinerja seluruh divisi dari satu dashboard')
bullet('Menyediakan arsip dokumen digital yang terorganisir per divisi')
bullet('Mengotomasi pengingat (reminder) melalui Telegram Bot')
bullet('Menyediakan fitur komunikasi internal berupa papan pengumuman')

heading('1.3 Pengguna Sistem', level=2)
add_table(['Role', 'Jumlah', 'Akses'], [
    ['Staff', '~15-20 orang', 'Input laporan, arsip divisi sendiri, ajukan izin'],
    ['Direksi/Pimpinan', '2 orang', 'Dashboard semua divisi, pengumuman, kelola user'],
])

doc.add_page_break()

# === 2. ARSITEKTUR SISTEM ===
heading('2. Arsitektur Sistem')
para('Sistem dibangun dengan arsitektur modern berbasis cloud yang seluruhnya menggunakan layanan gratis (free tier):')

add_table(['Komponen', 'Teknologi', 'Fungsi'], [
    ['Frontend + Backend', 'Next.js 15 (App Router)', 'Render halaman, form submit, API routes'],
    ['Database', 'Supabase PostgreSQL', 'Penyimpanan semua data (500 MB free)'],
    ['Authentication', 'Supabase Auth', 'Login, session, role management'],
    ['File Storage', 'Supabase Storage', 'Foto bukti & arsip dokumen (1 GB free)'],
    ['Hosting', 'Vercel (Free)', 'Deployment dan CDN'],
    ['Notifikasi', 'Telegram Bot API', 'Reminder harian & weekly digest'],
    ['UI Components', 'shadcn/ui + Tailwind CSS', 'Komponen antarmuka modern'],
    ['Domain', 'laporan.mahiratour.id', 'Subdomain dari domain yang sudah ada'],
])

para('')
para('Biaya Operasional Bulanan: Rp 0 (Nol Rupiah)', bold=True)
para('Seluruh layanan menggunakan free tier. Domain sudah dimiliki oleh perusahaan.')

doc.add_page_break()

# === 3. DAFTAR FITUR ===
heading('3. Daftar Fitur Lengkap')

heading('3.1 Fitur CORE (Utama)', level=2)
add_table(['No', 'Fitur', 'Deskripsi'], [
    ['1', 'Auth + Role', 'Login email/password. 2 role: Staff & Direksi'],
    ['2', 'Manajemen User & Divisi', 'CRUD user, assign ke divisi'],
    ['3', 'Laporan Harian', 'Pagi: input task. Sore: update status + submit'],
    ['4', 'Dashboard Divisi', 'Tiap divisi punya dashboard sendiri'],
    ['5', 'Dashboard Pimpinan', 'Overview semua divisi + drill-down'],
    ['6', 'Notifikasi Telegram', 'Reminder jam 16:00 "Jangan lupa submit laporan"'],
    ['7', 'Search & Filter', 'Cari laporan by tanggal, divisi, status, keyword'],
    ['8', 'Izin/Ketidakhadiran', 'Staff tandai izin, sistem tahu siapa yang tidak submit'],
])

heading('3.2 Fitur SHOULD HAVE (Penting)', level=2)
add_table(['No', 'Fitur', 'Deskripsi'], [
    ['9', 'Arsip Dokumen Divisi', 'Upload, browse, download file per divisi'],
    ['10', 'Papan Pengumuman', 'Pimpinan broadcast info ke semua/per divisi'],
    ['11', 'Catatan Pimpinan', 'Pimpinan beri feedback di laporan staff'],
    ['12', 'Upload Bukti Foto', 'Lampirkan foto di laporan harian'],
])

heading('3.3 Fitur NICE TO HAVE (Tambahan)', level=2)
add_table(['No', 'Fitur', 'Deskripsi'], [
    ['13', 'Weekly Digest Telegram', 'Ringkasan mingguan otomatis ke pimpinan'],
    ['14', 'Pin Dokumen Penting', 'Tandai dokumen penting agar selalu di atas'],
    ['15', 'Laporan Bulanan', 'Auto-recap dari data harian + narasi manual'],
])

doc.add_page_break()

# === 4. ALUR KERJA ===
heading('4. Alur Kerja Sistem')

heading('4.1 Alur Kerja Harian', level=2)

para('PAGI (07:00 - 09:00) - Staff Mulai Hari', bold=True)
bullet('Staff buka web app di HP')
bullet('Cek pengumuman terbaru dari pimpinan')
bullet('Input rencana kerja hari ini (task + prioritas)')
bullet('Task kemarin yang belum selesai otomatis muncul (carry-forward)')
bullet('Jika izin: klik "Izin Hari Ini", pilih alasan')

para('SIANG - Staff Bekerja Normal', bold=True)
bullet('Staff bekerja seperti biasa, tidak perlu buka sistem')
bullet('Jika ada dokumen penting: upload ke arsip divisi')

para('SORE (16:00 - 17:00) - Submit Laporan', bold=True)
bullet('Jam 16:00: Telegram Bot kirim reminder otomatis')
bullet('Staff buka web app, update status setiap task')
bullet('Tambah task baru jika ada kerjaan mendadak')
bullet('Upload foto bukti (opsional)')
bullet('Klik "Submit Laporan"')

para('MALAM - Pimpinan Review', bold=True)
bullet('Pimpinan buka dashboard: lihat siapa sudah submit, siapa belum')
bullet('Drill-down ke divisi tertentu untuk lihat detail')
bullet('Beri catatan/feedback di laporan staff')
bullet('Post pengumuman jika perlu')

heading('4.2 Alur Kerja Mingguan', level=2)
bullet('Setiap Senin pagi: sistem kirim Weekly Digest via Telegram ke pimpinan')
bullet('Berisi ringkasan: berapa hari submit, completion rate, task carry-forward')

heading('4.3 Alur Kerja Bulanan', level=2)
bullet('Akhir bulan: sistem auto-generate rekap bulanan per divisi')
bullet('Staff tambah narasi: capaian, kendala, rencana bulan depan')
bullet('Pimpinan review & acknowledge')

doc.add_page_break()

# === 5. AKSES KONTROL ===
heading('5. Hak Akses Pengguna')
add_table(['Data/Fitur', 'Staff', 'Pimpinan'], [
    ['Laporan sendiri', 'Baca & tulis', 'Baca semua + catatan'],
    ['Laporan divisi lain', 'Tidak bisa', 'Baca semua'],
    ['Arsip divisi sendiri', 'Upload & baca', 'Baca & download'],
    ['Arsip divisi lain', 'Tidak bisa', 'Baca & download'],
    ['Pengumuman', 'Baca saja', 'Baca & buat'],
    ['Izin', 'Ajukan sendiri', 'Lihat semua'],
    ['Kelola user', 'Tidak bisa', 'CRUD (Tambah/Edit/Hapus)'],
])

doc.add_page_break()

# === 6. TAMPILAN MENU ===
heading('6. Struktur Menu Aplikasi')

heading('6.1 Menu Staff', level=2)
add_table(['Menu', 'Isi'], [
    ['Beranda', 'Pengumuman terbaru + status laporan hari ini'],
    ['Laporan Saya', 'Input task pagi, update sore, histori laporan'],
    ['Arsip Divisi', 'Upload & browse dokumen divisi sendiri'],
    ['Izin', 'Ajukan izin hari ini'],
    ['Profil', 'Data diri, Telegram ID'],
])

heading('6.2 Menu Pimpinan', level=2)
add_table(['Menu', 'Isi'], [
    ['Dashboard', 'Overview semua divisi (submit/belum/izin)'],
    ['Per Divisi', 'Drill-down: laporan + arsip + statistik'],
    ['Pengumuman', 'Buat & kelola pengumuman'],
    ['Arsip Semua', 'Browse & download dokumen dari semua divisi'],
    ['Kelola User', 'Tambah/edit staff & divisi'],
    ['Laporan Bulanan', 'Review rekap bulanan semua divisi'],
])

doc.add_page_break()

# === 7. DATABASE ===
heading('7. Struktur Database')
para('Sistem menggunakan 13 tabel database yang saling terhubung:')

add_table(['No', 'Nama Tabel', 'Fungsi'], [
    ['1', 'divisions', 'Data divisi perusahaan'],
    ['2', 'users', 'Data karyawan + role + divisi'],
    ['3', 'daily_work_plans', 'Rencana kerja harian (input pagi)'],
    ['4', 'plan_tasks', 'Detail task dalam rencana kerja'],
    ['5', 'daily_reports', 'Laporan harian (submit sore)'],
    ['6', 'task_updates', 'Update status setiap task'],
    ['7', 'report_attachments', 'Foto bukti lampiran laporan'],
    ['8', 'report_notes', 'Catatan/feedback pimpinan'],
    ['9', 'absences', 'Data izin/ketidakhadiran'],
    ['10', 'division_documents', 'Arsip dokumen per divisi'],
    ['11', 'announcements', 'Data pengumuman'],
    ['12', 'announcement_reads', 'Status baca pengumuman'],
    ['13', 'monthly_reports', 'Rekap laporan bulanan'],
])

doc.add_page_break()

# === 8. TIMELINE ===
heading('8. Timeline Development (4 Minggu)')

heading('Minggu 1: Fondasi', level=2)
add_table(['Hari', 'Task', 'Output'], [
    ['1', 'Setup proyek + deploy', 'Halaman live di laporan.mahiratour.id'],
    ['2-3', 'Auth (login/logout) + role guard', 'Staff & direksi bisa login'],
    ['4', 'CRUD User & Divisi', 'Direksi bisa kelola staff & divisi'],
    ['5', 'Halaman izin', 'Staff bisa ajukan izin'],
])

heading('Minggu 2: Core Reporting', level=2)
add_table(['Hari', 'Task', 'Output'], [
    ['1-2', 'Input task pagi', 'Staff bisa input rencana kerja'],
    ['3-4', 'Update status sore + submit', 'Staff bisa submit laporan'],
    ['5', 'Carry-forward + upload foto', 'Task kemarin otomatis muncul'],
])

heading('Minggu 3: Dashboard & Komunikasi', level=2)
add_table(['Hari', 'Task', 'Output'], [
    ['1-2', 'Dashboard divisi', 'Staff lihat statistik divisi'],
    ['3-4', 'Dashboard pimpinan', 'Direksi lihat overview + detail'],
    ['5', 'Pengumuman', 'Direksi bisa broadcast pengumuman'],
])

heading('Minggu 4: Dokumen & Polish', level=2)
add_table(['Hari', 'Task', 'Output'], [
    ['1-2', 'Arsip dokumen divisi', 'Upload, browse, download, pin'],
    ['3', 'Catatan pimpinan + Telegram', 'Feedback + notifikasi jalan'],
    ['4', 'Testing & bug fix', 'Semua flow berjalan lancar'],
    ['5', 'UAT + training', 'User pertama mulai pakai'],
])

doc.add_page_break()

# === 9. BIAYA ===
heading('9. Estimasi Biaya')
add_table(['Komponen', 'Biaya/Bulan', 'Keterangan'], [
    ['Hosting (Vercel)', 'Rp 0', 'Free tier, cukup untuk 15-20 user'],
    ['Database (Supabase)', 'Rp 0', 'Free tier: 500MB database, 1GB storage'],
    ['Domain', 'Rp 0', 'Subdomain dari mahiratour.id yang sudah ada'],
    ['Telegram Bot', 'Rp 0', 'Gratis dari Telegram'],
    ['Cron Job', 'Rp 0', 'cron-job.org free tier'],
    ['TOTAL', 'Rp 0/bulan', 'Bisa diupgrade jika perlu di masa depan'],
])

doc.add_paragraph()
para('Catatan: Jika di masa depan user bertambah atau storage penuh, biaya upgrade Supabase mulai dari ~$25/bulan (~Rp 400.000). Namun untuk kebutuhan saat ini, free tier sudah sangat mencukupi.')

doc.add_page_break()

# === 10. PENUTUP ===
heading('10. Penutup')
para('Blueprint ini merupakan dokumen teknis yang menggambarkan secara lengkap rancangan Sistem Laporan Terintegrasi untuk Mahira Tour. Sistem ini dirancang untuk:')
bullet('Meningkatkan efisiensi pelaporan kerja harian')
bullet('Mempermudah pimpinan dalam memonitor kinerja seluruh divisi')
bullet('Mendigitalkan arsip dokumen perusahaan')
bullet('Mengotomasi pengingat dan komunikasi internal')

para('')
para('Dengan biaya operasional Rp 0/bulan dan timeline development 4 minggu, sistem ini diharapkan dapat memberikan dampak positif yang signifikan terhadap produktivitas dan efisiensi kerja di Mahira Tour.')

doc.add_paragraph()
doc.add_paragraph()
para('Sungai Penuh, 28 April 2026', align=WD_ALIGN_PARAGRAPH.RIGHT)
para('Disusun oleh,', align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_paragraph()
doc.add_paragraph()
para('Khoirul Gunawan', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
para('Staff IT / FullStack Developer', align=WD_ALIGN_PARAGRAPH.RIGHT)

fp = os.path.join(OUT, 'BLUEPRINT SISTEM LAPORAN TERINTEGRASI.docx')
doc.save(fp)
print(f'[OK] {fp}')
